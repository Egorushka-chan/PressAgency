from tkinter import *
from tkinter import ttk

import DBAccessor
from model.user import UserRank
import openpyxl as xl

import datetime
import docx
import csv

# https://habr.com/ru/company/skillfactory/blog/553224/

class ReportForm:
    def __init__(self, user):
        self.root = Tk()
        self.root.title('Отчеты')
        self.user = user
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

        self.reportsFrame = Frame(self.root)
        self.reportsBox = Listbox(self.reportsFrame)
        self.reportScroll = Scrollbar(self.reportsFrame, orient="vertical", command=self.reportsBox.yview)

        self.reportsBox.configure(yscrollcommand=self.reportScroll.set)
        self.reportsBox.bind('<<ListboxSelect>>', self.report_select)

        self.reportsBox.pack(side='right')
        self.reportScroll.pack(side='left', fill='y')

        self.reportsFrame.grid(row=0, column=0)

        self.reports = DBAccessor.select_info(fr0m='Reports')
        for report in self.reports:
            self.reportsBox.insert(report[0], f'{report[0]} {report[1].strip("{}")}')

        if user.rank in (UserRank.employee, UserRank.admin):
            Button(self.root, text='Добавить').grid(row=1, column=0)
            self.changeButton = Button(self.root, text='Изменить', state=DISABLED)
            self.changeButton.grid(row=2, column=0)
            if user.rank is UserRank.admin:
                self.deleteButton = Button(self.root, text='Удалить', state=DISABLED)
                self.deleteButton.grid(row=3, column=0)
                # TODO: Все это нужно сделать полезным (просто сделать)

        self.tableReportsFrame = Frame(self.root)
        self.tableQuery = ttk.Treeview(self.tableReportsFrame, height=10)

        self.tableQuery.pack(side='right')
        self.tableScroll = ttk.Scrollbar(self.tableReportsFrame, orient="vertical", command=self.tableQuery.yview)
        self.tableScroll.pack(side='left', fill='y')
        self.tableQuery.configure(yscrollcommand=self.tableScroll.set)

        self.tableReportsFrame.grid(row=0, column=1, rowspan=4, padx=5, pady=5)

        self.rightFrame = Frame(self.root)
        self.wordButton = Button(self.rightFrame, text='Вывести в Word', command=self.word_output, state=DISABLED)
        self.wordButton.pack()
        self.csvButton = Button(self.rightFrame, text='Вывести в CSV', command=self.csv_output, state=DISABLED)
        self.csvButton.pack()
        self.excelButton = Button(self.rightFrame, text='Вывести в Excel', command=self.excel_output, state=DISABLED)
        self.excelButton.pack()
        self.rightFrame.grid(row=0, column=2, rowspan=4)

        self.data = None
        self.data_columns_names = None

        self.root.mainloop()





    def report_select(self, trash=1):
        if self.user.rank in (UserRank.employee, UserRank.admin):
            self.changeButton.configure(state='normal')
            if self.user.rank is UserRank.admin:
                self.deleteButton.configure(state='normal')

        selected_value = self.reportsBox.selection_get()
        id = selected_value.split(' ')[0]
        required_report = [report for report in self.reports if str(report[0]) == id]
        raw_query = required_report[0][2]

        processed_query = DBAccessor.select_info(raw_query=raw_query, keys=True)

        self.data = processed_query[0]
        self.data_columns_names = processed_query[1]

        self.tableQuery.delete(*self.tableQuery.get_children())

        tab_read_columns = []
        for i in range(1, len(self.data_columns_names) + 1):
            tab_read_columns.append(str(i))

        self.tableQuery.configure(show='headings', columns=tab_read_columns)
        for col in tab_read_columns:
            self.tableQuery.heading(col, text=col, command=lambda _col=col: \
                self.treeview_sort_column(self.tableQuery, _col, False))
        for i, column_name in enumerate(self.data_columns_names):
            self.tableQuery.column(str(i + 1), anchor='c')
            self.tableQuery.heading(str(i + 1), text=column_name)

        for row in self.data:
            self.tableQuery.insert("", 'end', values=row)

        self.wordButton.configure(state='normal')
        self.csvButton.configure(state='normal')
        self.excelButton.configure(state='normal')

    def word_output(self):
        # values = self.tableQuery.item(*self.tableQuery.get_children(), option='values')
        document = docx.Document()

        document.add_paragraph('Report (Отчет по русски)')

        table = document.add_table(rows=1, cols=0)
        for i, column in enumerate(self.data_columns_names):
            table.add_column(len(column) * 12)
            upper_row = table.rows[0].cells
            upper_row[i].text = column

        for row in self.data:
            current_row = table.add_row().cells
            for i, value in enumerate(row):
                current_row[i].text = value

        document.save(
            DBAccessor.base_path + f'reports/{self.reportsBox.selection_get()}_{datetime.datetime.now().date()}.docx')
        # TODO: таблица так не работает. Сделать через Template

    def csv_output(self):
        file_name = DBAccessor.base_path + f'reports/{self.reportsBox.selection_get()}_{datetime.datetime.now().date()}.csv'
        with open(file_name, 'w', encoding='UTF-8', newline='') as file:
            writer = csv.writer(file, delimiter=';')
            writer.writerow(self.data_columns_names)

        for row in self.data:
            with open(file_name, 'a', encoding='UTF-8', newline='') as file:
                writer = csv.writer(file, delimiter=';')
                writer.writerow(row)

    def excel_output(self):
        workbook = xl.Workbook()
        sheet = workbook.active
        sheet.append(self.data_columns_names)
        sheet.title = "Результат"
        for row in self.data:
            sheet.append(row)
        workbook.save(DBAccessor.base_path + f'reports/{self.reportsBox.selection_get()}_{datetime.datetime.now().date()}.xlsx')

    def treeview_sort_column(self, tv, col, reverse):
        l = [(tv.set(k, col), k) for k in tv.get_children('')]
        l.sort(reverse=reverse)

        # rearrange items in sorted positions
        for index, (val, k) in enumerate(l):
            tv.move(k, '', index)

        # reverse sort next time
        tv.heading(col, command=lambda: \
            self.treeview_sort_column(tv, col, not reverse))

    def on_closing(self):
        self.root.destroy()
